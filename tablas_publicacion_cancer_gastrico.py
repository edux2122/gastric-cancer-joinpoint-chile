# =============================================================================
# TABLAS DEFINITIVAS PARA PUBLICACIÓN CIENTÍFICA
# Cáncer Gástrico en Chile, 2018-2022
# Tabla 1: Tasas crudas | Tabla 2: Tasas ajustadas por edad
#
# Equivalente Python/Colab del script R con flextable/officer
# Replica el estilo booktabs (Arial 9pt, bordes solo top/bottom, celdas fusionadas)
# Lee numeradores desde Excel de preprocesamiento y poblaciones desde CSV
# =============================================================================

# --- INSTALACIÓN (ejecutar una sola vez en Colab) ---
# !pip install python-docx pandas openpyxl --quiet

import pandas as pd
import numpy as np
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from google.colab import drive

drive.mount('/content/drive')

# =============================================================================
# 1. RUTAS
# =============================================================================

base_dir    = '/content/drive/MyDrive/INVESTIGACIONES/Descripcion epidemiológica del cáncer gástrico en Chile/'
excel_path  = base_dir + 'Datos egresos/tabla_egreso_cancer_gastrico_quinquenios.xlsx'
bm_path     = base_dir + 'Datos egresos/BANCOMUNDIAL.csv'
censo_path  = base_dir + 'Datos egresos/CENSO2024.csv'
output_file = base_dir + 'Tablas_Finales_Publicacion.docx'

AÑOS = [2018, 2019, 2020, 2021, 2022]

# =============================================================================
# 2. CARGA DE POBLACIONES
# =============================================================================

# --- Banco Mundial ---
# Columnas: POBLACION, 2018, 2019, 2020, 2021, 2022
# Filas (índice): 0=hombres total, 1=mujeres total, 2=general total,
#                 3=0-14H, 4=0-14M, 5=0-14G, 6=15-64H, 7=15-64M, 8=15-64G,
#                 9=65+H, 10=65+M, 11=65+G
bm_raw = pd.read_csv(bm_path, encoding='latin1')
bm_raw.columns = ['indicador'] + [str(a) for a in AÑOS]

mapping_bm = {
    ('GENERAL', '0-14'):  5,  ('GENERAL', '15-64'): 8,  ('GENERAL', '65+'): 11,
    ('HOMBRE',  '0-14'):  3,  ('HOMBRE',  '15-64'): 6,  ('HOMBRE',  '65+'):  9,
    ('MUJER',   '0-14'):  4,  ('MUJER',   '15-64'): 7,  ('MUJER',   '65+'): 10,
}
pob_bm = {
    k: {int(a): int(bm_raw.iloc[r][str(a)]) for a in AÑOS}
    for k, r in mapping_bm.items()
}

# --- Censo 2024 (estándar para estandarización directa) ---
# Columnas: EDAD, SEXO, N | EDAD: 0-14 / 15-64 / 65+ | SEXO: Hombres / Mujeres
censo_raw = pd.read_csv(censo_path, encoding='latin1')
sexo_censo = {'Hombres': 'HOMBRE', 'Mujeres': 'MUJER'}
pob_censo = {}
for _, row in censo_raw.iterrows():
    pob_censo[(sexo_censo[row['SEXO'].strip()], row['EDAD'].strip())] = int(row['N'])
for edad in ['0-14', '15-64', '65+']:
    pob_censo[('GENERAL', edad)] = pob_censo[('HOMBRE', edad)] + pob_censo[('MUJER', edad)]

# =============================================================================
# 3. CARGA DEL EXCEL DE EGRESOS Y CÁLCULO DE TASAS
# =============================================================================

df_eg   = pd.read_excel(excel_path, engine='openpyxl')
edad_df = df_eg[df_eg['VARIABLE'] == 'GRUPO_EDAD_QUINQUENIO'].copy()

grupos_q = {
    '0-14':  ['00-04 años', '05-09 años', '10-14 años'],
    '15-64': ['15-19 años', '20-24 años', '25-29 años', '30-34 años', '35-39 años',
              '40-44 años', '45-49 años', '50-54 años', '55-59 años', '60-64 años'],
    '65+':   ['65-69 años', '70-74 años', '75-79 años', '80-84 años', '85 y más años'],
}
etiq_grp = {'0-14': '0-14 años', '15-64': '15-64 años', '65+': '65 y más años'}
sexo_lbl = {'GENERAL': 'General', 'HOMBRE': 'Hombre', 'MUJER': 'Mujer'}

# --- Tabla 1: Tasas crudas ---
rows_t1 = []
for sr, sl in sexo_lbl.items():
    for gk, cats in grupos_q.items():
        sub   = edad_df[(edad_df['SEXO'] == sr) & (edad_df['CATEGORIA'].isin(cats))]
        n_tot = sub['TOTAL_5A'].fillna(0).sum()
        p_tot = sum(pob_bm[(sr, gk)][a] for a in AÑOS)
        row   = {'sexo': sl, 'grupo_edad': etiq_grp[gk],
                 'tasa_periodo': round((n_tot / p_tot) * 100000, 2)}
        for a in AÑOS:
            n_a = sub[str(a)].fillna(0).sum()
            row[str(a)] = round((n_a / pob_bm[(sr, gk)][a]) * 100000, 2)
        rows_t1.append(row)
df_t1 = pd.DataFrame(rows_t1)

# --- Tabla 2: Tasas ajustadas (estandarización directa, estándar Censo 2024) ---
rows_t2 = []
for sr, sl in sexo_lbl.items():
    for a in AÑOS:
        num = den = 0.0
        for gk, cats in grupos_q.items():
            sub  = edad_df[(edad_df['SEXO'] == sr) & (edad_df['CATEGORIA'].isin(cats))]
            n_a  = sub[str(a)].fillna(0).sum()
            p_a  = pob_bm[(sr, gk)][a]
            te   = (n_a / p_a) * 100000 if p_a > 0 else 0
            ps   = pob_censo[(sr, gk)]
            num += te * ps
            den += ps
        rows_t2.append({'sexo': sl, 'anio': a, 'tasa_ajustada': round(num / den, 2)})
df_t2 = pd.DataFrame(rows_t2)

print("=== TABLA 1 — TASAS CRUDAS ===")
print(df_t1.to_string(index=False))
print("\n=== TABLA 2 — TASAS AJUSTADAS ===")
print(df_t2.to_string(index=False))

# =============================================================================
# 4. HELPERS — XML exacto del docx generado por R/flextable
# Estilo booktabs: Arial 9pt (sz=18), spacing before/after=80, ind=80
# Bordes: top/bot single sz=12 en header; bot sz=12 en última fila de cada sexo
# Sin bordes laterales | vMerge para columna Sexo fusionada
# =============================================================================

def make_tcPr(top_val='none', top_sz='0',
              bot_val='none', bot_sz='0',
              vmerge=None, valign='center'):
    tcPr = OxmlElement('w:tcPr')
    if vmerge == 'restart':
        vm = OxmlElement('w:vMerge'); vm.set(qn('w:val'), 'restart'); tcPr.append(vm)
    elif vmerge == 'continue':
        vm = OxmlElement('w:vMerge'); tcPr.append(vm)
    borders = OxmlElement('w:tcBorders')
    for side, val, sz in [('bottom', bot_val, bot_sz), ('top', top_val, top_sz),
                           ('left',  'none',  '0'),    ('right', 'none', '0')]:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), val); el.set(qn('w:sz'), sz)
        el.set(qn('w:space'), '0'); el.set(qn('w:color'), '000000')
        borders.append(el)
    tcPr.append(borders)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'FFFFFF'); tcPr.append(shd)
    mar = OxmlElement('w:tcMar')
    for side in ['top', 'bottom', 'left', 'right']:
        m = OxmlElement(f'w:{side}')
        m.set(qn('w:w'), '0'); m.set(qn('w:type'), 'dxa'); mar.append(m)
    tcPr.append(mar)
    va = OxmlElement('w:vAlign'); va.set(qn('w:val'), valign); tcPr.append(va)
    return tcPr


def set_cell(cell, text, align='left',
             top_val='none', top_sz='0',
             bot_val='none', bot_sz='0',
             vmerge=None, valign='center'):
    tc  = cell._tc
    old = tc.find(qn('w:tcPr'))
    if old is not None: tc.remove(old)
    tc.insert(0, make_tcPr(top_val, top_sz, bot_val, bot_sz, vmerge, valign))
    for p in cell.paragraphs:
        p._p.getparent().remove(p._p)
    p_el = OxmlElement('w:p')
    pPr  = OxmlElement('w:pPr')
    ps   = OxmlElement('w:pStyle'); ps.set(qn('w:val'), 'Normal'); pPr.append(ps)
    jc   = OxmlElement('w:jc'); jc.set(qn('w:val'), align); pPr.append(jc)
    pBdr = OxmlElement('w:pBdr')
    for side in ['bottom', 'top', 'left', 'right']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'none'); b.set(qn('w:sz'), '0')
        b.set(qn('w:space'), '0'); b.set(qn('w:color'), '000000')
        pBdr.append(b)
    pPr.append(pBdr)
    sp = OxmlElement('w:spacing')
    sp.set(qn('w:after'), '80'); sp.set(qn('w:before'), '80')
    sp.set(qn('w:line'), '240'); pPr.append(sp)
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), '80'); ind.set(qn('w:right'), '80')
    ind.set(qn('w:firstLine'), '0'); ind.set(qn('w:firstLineChars'), '0')
    pPr.append(ind)
    rPr_p = OxmlElement('w:rPr')
    fonts = OxmlElement('w:rFonts')
    for attr in ['w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs']:
        fonts.set(qn(attr), 'Arial')
    rPr_p.append(fonts)
    for tag, val in [('w:i', 'false'), ('w:b', 'false'),
                     ('w:u', 'none'),  ('w:strike', 'false')]:
        el = OxmlElement(tag); el.set(qn('w:val'), val); rPr_p.append(el)
    for tag, v in [('w:sz', '18'), ('w:szCs', '18')]:
        el = OxmlElement(tag); el.set(qn('w:val'), v); rPr_p.append(el)
    col = OxmlElement('w:color'); col.set(qn('w:val'), '000000'); rPr_p.append(col)
    pPr.append(rPr_p)
    p_el.append(pPr)
    r_el = OxmlElement('w:r')
    rPr  = OxmlElement('w:rPr')
    fonts2 = OxmlElement('w:rFonts')
    for attr in ['w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs']:
        fonts2.set(qn(attr), 'Arial')
    rPr.append(fonts2)
    for tag, val in [('w:i', 'false'), ('w:b', 'false'),
                     ('w:u', 'none'),  ('w:strike', 'false')]:
        el = OxmlElement(tag); el.set(qn('w:val'), val); rPr.append(el)
    for tag, v in [('w:sz', '18'), ('w:szCs', '18')]:
        el = OxmlElement(tag); el.set(qn('w:val'), v); rPr.append(el)
    col2 = OxmlElement('w:color'); col2.set(qn('w:val'), '000000'); rPr.append(col2)
    r_el.append(rPr)
    t = OxmlElement('w:t')
    if text and (str(text)[0] == ' ' or str(text)[-1] == ' '):
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t.text = str(text)
    r_el.append(t); p_el.append(r_el); tc.append(p_el)


def set_tblGrid(tbl, widths_dxa):
    tblEl = tbl._tbl
    old   = tblEl.find(qn('w:tblGrid'))
    if old is not None: tblEl.remove(old)
    tblGrid = OxmlElement('w:tblGrid')
    for w in widths_dxa:
        gc = OxmlElement('w:gridCol'); gc.set(qn('w:w'), str(w)); tblGrid.append(gc)
    tblPr = tblEl.find(qn('w:tblPr'))
    tblEl.insert(list(tblEl).index(tblPr) + 1, tblGrid)


def fmt(val, dec=2):
    if val is None or (isinstance(val, float) and np.isnan(val)):
        return '—'
    return f"{val:.{dec}f}".replace('.', ',')


def add_caption(doc, text):
    p   = doc.add_paragraph(); p._p.clear()
    pPr = OxmlElement('w:pPr')
    ps  = OxmlElement('w:pStyle'); ps.set(qn('w:val'), 'Normal'); pPr.append(ps)
    sp  = OxmlElement('w:spacing')
    sp.set(qn('w:after'), '80'); sp.set(qn('w:before'), '80'); pPr.append(sp)
    p._p.append(pPr)
    r   = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    f   = OxmlElement('w:rFonts')
    for a in ['w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs']: f.set(qn(a), 'Arial')
    rPr.append(f)
    for tag, v in [('w:sz', '18'), ('w:szCs', '18')]:
        el = OxmlElement(tag); el.set(qn('w:val'), v); rPr.append(el)
    r.append(rPr)
    t = OxmlElement('w:t'); t.text = text; r.append(t); p._p.append(r)


def add_footnote(doc, text):
    p   = doc.add_paragraph(); p._p.clear()
    pPr = OxmlElement('w:pPr')
    ps  = OxmlElement('w:pStyle'); ps.set(qn('w:val'), 'Normal'); pPr.append(ps)
    sp  = OxmlElement('w:spacing')
    sp.set(qn('w:after'), '40'); sp.set(qn('w:before'), '40'); pPr.append(sp)
    p._p.append(pPr)
    r   = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    f   = OxmlElement('w:rFonts')
    for a in ['w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs']: f.set(qn(a), 'Arial')
    rPr.append(f)
    i   = OxmlElement('w:i'); i.set(qn('w:val'), 'true'); rPr.append(i)
    for tag, v in [('w:sz', '16'), ('w:szCs', '16')]:
        el = OxmlElement(tag); el.set(qn('w:val'), v); rPr.append(el)
    r.append(rPr)
    t = OxmlElement('w:t'); t.text = text; r.append(t); p._p.append(r)

# =============================================================================
# 5. CONSTRUCCIÓN TABLA 1
# tblGrid replicado del docx R: [1872, 1872, 1296, 1296, 1296, 1296, 1296, 1296]
# =============================================================================

def build_tabla1(doc, df):
    GRID    = [1872, 1872, 1296, 1296, 1296, 1296, 1296, 1296]
    orden_s = ['General', 'Hombre', 'Mujer']
    orden_e = ['0-14 años', '15-64 años', '65 y más años']

    tbl = doc.add_table(rows=0, cols=8)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tblPr = tbl._tbl.find(qn('w:tblPr'))
    ts = tblPr.find(qn('w:tblStyle'))
    if ts is not None: tblPr.remove(ts)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '0'); tblW.set(qn('w:type'), 'auto')
    tblPr.append(tblW)
    set_tblGrid(tbl, GRID)

    # Header
    hdr    = tbl.add_row()
    labels = ['Sexo', 'Grupo de edad', 'Total periodo\n2018-2022',
              '2018', '2019', '2020', '2021', '2022']
    aligns = ['left', 'left', 'right', 'right', 'right', 'right', 'right', 'right']
    for cell, lbl, aln in zip(hdr.cells, labels, aligns):
        set_cell(cell, lbl, align=aln,
                 top_val='single', top_sz='12',
                 bot_val='single', bot_sz='12')

    # Datos
    for s_i, sexo in enumerate(orden_s):
        sub_s = df[df['sexo'] == sexo]
        for e_i, edad in enumerate(orden_e):
            rd    = sub_s[sub_s['grupo_edad'] == edad].iloc[0]
            last  = (e_i == len(orden_e) - 1)
            bot_v = 'single' if last else 'none'
            bot_s = '12'     if last else '0'
            fila  = tbl.add_row()

            set_cell(fila.cells[0], sexo if e_i == 0 else '',
                     align='left', top_val='none', top_sz='0',
                     bot_val='none', bot_sz='0',
                     vmerge='restart' if e_i == 0 else 'continue',
                     valign='top')

            set_cell(fila.cells[1], edad, align='left',
                     top_val='none', top_sz='0',
                     bot_val=bot_v, bot_sz=bot_s)

            vals = [fmt(rd['tasa_periodo'])] + [fmt(rd[str(a)]) for a in AÑOS]
            for c_i, val in enumerate(vals, start=2):
                set_cell(fila.cells[c_i], val, align='right',
                         top_val='none', top_sz='0',
                         bot_val=bot_v, bot_sz=bot_s)

    return tbl

# =============================================================================
# 6. CONSTRUCCIÓN TABLA 2
# =============================================================================

def build_tabla2(doc, df):
    GRID    = [2160, 2160, 2160]
    orden_s = ['General', 'Hombre', 'Mujer']

    tbl = doc.add_table(rows=0, cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tblPr = tbl._tbl.find(qn('w:tblPr'))
    ts = tblPr.find(qn('w:tblStyle'))
    if ts is not None: tblPr.remove(ts)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '0'); tblW.set(qn('w:type'), 'auto')
    tblPr.append(tblW)
    set_tblGrid(tbl, GRID)

    # Header
    hdr = tbl.add_row()
    for cell, lbl, aln in zip(hdr.cells,
                               ['Sexo', 'Año', 'Tasa ajustada*'],
                               ['left', 'right', 'right']):
        set_cell(cell, lbl, align=aln,
                 top_val='single', top_sz='12',
                 bot_val='single', bot_sz='12')

    # Datos
    for s_i, sexo in enumerate(orden_s):
        sub_s = df[df['sexo'] == sexo].sort_values('anio')
        n_a   = len(sub_s)
        for a_i, (_, row) in enumerate(sub_s.iterrows()):
            last  = (a_i == n_a - 1)
            bot_v = 'single' if last else 'none'
            bot_s = '12'     if last else '0'
            fila  = tbl.add_row()

            set_cell(fila.cells[0], sexo if a_i == 0 else '',
                     align='left', top_val='none', top_sz='0',
                     bot_val='none', bot_sz='0',
                     vmerge='restart' if a_i == 0 else 'continue',
                     valign='top')

            set_cell(fila.cells[1], str(int(row['anio'])), align='right',
                     top_val='none', top_sz='0', bot_val=bot_v, bot_sz=bot_s)

            set_cell(fila.cells[2], fmt(row['tasa_ajustada']), align='right',
                     top_val='none', top_sz='0', bot_val=bot_v, bot_sz=bot_s)

    return tbl

# =============================================================================
# 7. GENERACIÓN DEL DOCUMENTO
# =============================================================================

doc = Document()
sec = doc.sections[0]
sec.top_margin    = Inches(1)
sec.bottom_margin = Inches(1)
sec.left_margin   = Inches(1)
sec.right_margin  = Inches(1)

add_caption(doc,
    "Tabla 1. Tasas crudas de egresos hospitalarios por cáncer gástrico según "
    "grupo de edad, sexo y año. Chile, 2018-2022 (tasas por 100.000 hab).")
build_tabla1(doc, df_t1)

doc.add_page_break()

add_caption(doc,
    "Tabla 2. Tasas de egresos hospitalarios por cáncer gástrico ajustadas por "
    "edad según sexo y año. Chile, 2018-2022 (tasas por 100.000 hab).")
build_tabla2(doc, df_t2)
add_footnote(doc, "*Ajustada por método directo usando población estándar Censo 2024.")

doc.save(output_file)
print(f"\n¡Éxito! Archivo guardado en:\n{output_file}")
